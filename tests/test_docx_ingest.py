from __future__ import annotations

from io import BytesIO

from docx import Document

from app.services.ingest.docx_ingest import ingest_docx_bytes


def _build_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph("本项目为示例工程。")
    list_item = doc.add_paragraph("需提供近三年同类业绩。")
    list_item.style = "List Bullet"

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "工期"
    table.cell(1, 1).text = "120日历天"

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_ingest_docx_bytes_maps_heading_list_table_to_blocks() -> None:
    result = ingest_docx_bytes("demo.docx", _build_docx_bytes())

    block_types = [item.block_type for item in result.blocks]
    assert "TITLE" in block_types
    assert "PARA" in block_types
    assert "LIST" in block_types
    assert "TABLE" in block_types

    first_title = next(item for item in result.blocks if item.block_type == "TITLE")
    assert "第一章 总则" in first_title.content_text

    first_table = next(item for item in result.blocks if item.block_type == "TABLE")
    assert "指标" in first_table.content_text
    assert "120日历天" in first_table.content_text

    assert result.filename == "demo.docx"
    assert result.page_count == 1
