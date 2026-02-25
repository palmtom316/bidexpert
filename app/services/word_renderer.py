from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate

from app.core.config import settings

_HEADING_STYLES = {"Title1", "Title2", "Title3", "Title4"}
_PARAGRAPH_STYLES = {"BodyText", "BodyText_Indent", "ClauseText"}
_MANUAL_NUMBERING = re.compile(r"^\s*\d+(?:\.\d+)*[\.、\s]+")
_HAS_HTML = re.compile(r"<[^>]+>")
_HAS_MARKDOWN = re.compile(r"\*\*|__|`")
_HAS_STYLE_CODE = re.compile(r"\{\\[^}]+\}|w:(?:pPr|rPr|sectPr|br)")


_SIGNATURE_SECTIONS = {
    "法定代表人签章",
    "项目经理签章",
    "技术负责人签章",
    "投标人盖章",
    "授权代表签字",
    "日期",
}


def _add_signature_placeholder(doc: Document, label: str) -> None:
    """Add a signature placeholder line to the document."""
    paragraph = doc.add_paragraph(f"{label}：__________________")
    paragraph.paragraph_format.space_before = Cm(0.5)


def ensure_default_template(path: Path) -> None:
    if path.exists():
        return
    doc = Document()

    # Cover page
    doc.add_heading("投标文件", level=0)
    doc.add_paragraph("")
    doc.add_paragraph("项目名称：{{project_name}}")
    doc.add_paragraph("投标人：{{bidder_name}}")
    doc.add_paragraph("编制日期：{{compile_date}}")
    doc.add_paragraph("")

    # Standard sections
    doc.add_heading("目录", level=1)
    doc.add_paragraph("（由 Word 自动生成）")

    doc.add_heading("第一章 投标函", level=1)
    doc.add_paragraph("{{bid_letter}}")

    doc.add_heading("第二章 法定代表人身份证明", level=1)
    doc.add_paragraph("{{legal_representative}}")

    doc.add_heading("第三章 授权委托书", level=1)
    doc.add_paragraph("{{authorization}}")

    doc.add_heading("第四章 投标保证金", level=1)
    doc.add_paragraph("{{bid_bond}}")

    doc.add_heading("第五章 技术方案", level=1)
    doc.add_paragraph("{{technical_plan}}")

    doc.add_heading("第六章 施工组织设计", level=1)
    doc.add_paragraph("{{construction_plan}}")

    doc.add_heading("第七章 质量保证方案", level=1)
    doc.add_paragraph("{{quality_plan}}")

    doc.add_heading("第八章 安全生产方案", level=1)
    doc.add_paragraph("{{safety_plan}}")

    doc.add_heading("第九章 施工进度计划", level=1)
    doc.add_paragraph("{{schedule_plan}}")

    doc.add_heading("第十章 资源配置方案", level=1)
    doc.add_paragraph("{{resource_plan}}")

    doc.add_heading("第十一章 商务部分", level=1)
    doc.add_paragraph("{{commercial_proposal}}")

    # Headers and footers
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "{{project_name}} 投标文件"

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "{{bidder_name}} | 机密"

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _safe_path(raw_path: str, base_dir: Path) -> Path:
    candidate = Path(raw_path)
    full_path = candidate if candidate.is_absolute() else (base_dir / candidate)

    resolved_base = base_dir.resolve()
    resolved_target = full_path.resolve(strict=False)
    try:
        resolved_target.relative_to(resolved_base)
    except ValueError as exc:
        raise ValueError(f"path is outside allowed directory: {resolved_base}") from exc
    return resolved_target


def _apply_page_layout(doc: Document, style_config: dict[str, Any] | None) -> None:
    if not style_config:
        return
    page_cfg = style_config.get("page", {})
    if page_cfg:
        for section in doc.sections:
            if "marginTop" in page_cfg:
                section.top_margin = Cm(float(page_cfg["marginTop"]))
            if "marginBottom" in page_cfg:
                section.bottom_margin = Cm(float(page_cfg["marginBottom"]))
            if "marginLeft" in page_cfg:
                section.left_margin = Cm(float(page_cfg["marginLeft"]))
            if "marginRight" in page_cfg:
                section.right_margin = Cm(float(page_cfg["marginRight"]))
            if "headerOffset" in page_cfg:
                section.header_distance = Cm(float(page_cfg["headerOffset"]))
            if "footerOffset" in page_cfg:
                section.footer_distance = Cm(float(page_cfg["footerOffset"]))

    if style_config.get("generateTOC"):
        add_toc(doc)
        set_update_fields(doc)


def render_word(
    output_path: str,
    placeholders: dict[str, str],
    template_path: str | None = None,
    style_config: dict | None = None,
) -> str:
    template_root = Path(settings.render_template_dir)
    export_root = Path(settings.render_output_dir)
    template_root.mkdir(parents=True, exist_ok=True)
    export_root.mkdir(parents=True, exist_ok=True)

    template_candidate = template_path or "default_tender_template.docx"
    template_file = _safe_path(template_candidate, template_root)
    ensure_default_template(template_file)

    out = _safe_path(output_path, export_root)
    if out.suffix.lower() != ".docx":
        raise ValueError("output_path must end with .docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    tpl = DocxTemplate(str(template_file))
    tpl.render(placeholders)
    tpl.save(str(out))

    if style_config:
        doc = Document(str(out))
        _apply_page_layout(doc, style_config)
        doc.save(str(out))

    return str(out)


def add_toc(doc: Document) -> None:
    """
    Insert a Table of Contents field code.
    Field code: {TOC \\o "1-3" \\h \\z \\u}
    """
    paragraph = doc.add_paragraph()
    doc.element.body.insert(0, paragraph._element)

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def set_update_fields(doc: Document) -> None:
    settings_xml = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_xml.append(update_fields)


def _validate_text(value: str, *, is_heading: bool = False) -> str:
    text = value.strip()
    if not text:
        raise ValueError("text cannot be empty")
    if text == "目录" or text.lower() == "table of contents":
        raise ValueError("目录文本必须由 Word 自动目录生成")
    if "\f" in text or "[PAGE_BREAK]" in text or "<pagebreak" in text.lower():
        raise ValueError("禁止输出分页符标记")
    if _HAS_HTML.search(text):
        raise ValueError("禁止输出 HTML 标签")
    if _HAS_MARKDOWN.search(text):
        raise ValueError("禁止输出 Markdown 样式")
    if _HAS_STYLE_CODE.search(text):
        raise ValueError("禁止输出样式控制代码")
    if is_heading and _MANUAL_NUMBERING.match(text):
        raise ValueError("标题禁止手写编号")
    return text


def _ensure_paragraph_style(doc: Document, style_name: str, allowed: set[str]) -> None:
    if style_name not in allowed:
        raise ValueError(f"unsupported style: {style_name}")
    if style_name not in {style.name for style in doc.styles}:
        raise ValueError(f"template missing required style: {style_name}")


def _append_heading_block(doc: Document, block: dict[str, Any]) -> None:
    style_name = str(block.get("style") or "")
    _ensure_paragraph_style(doc, style_name, _HEADING_STYLES)
    text = _validate_text(str(block.get("text") or ""), is_heading=True)
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name


def _append_paragraph_block(doc: Document, block: dict[str, Any]) -> None:
    style_name = str(block.get("style") or "")
    _ensure_paragraph_style(doc, style_name, _PARAGRAPH_STYLES)
    text = _validate_text(str(block.get("text") or ""))
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name
    risk_level = str(block.get("risk_level") or "").strip().lower()
    if risk_level == "high" and paragraph.runs:
        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    evidence = block.get("evidence")
    if isinstance(evidence, list) and evidence:
        _ensure_paragraph_style(doc, "ClauseText", _PARAGRAPH_STYLES)
        evidence_rows: list[str] = []
        for idx, item in enumerate(evidence, start=1):
            if not isinstance(item, dict):
                continue
            doc_id = _validate_text(str(item.get("doc_id") or "unknown"))
            chunk_id = _validate_text(str(item.get("chunk_id") or "unknown"))
            page_range = item.get("page_range") if isinstance(item.get("page_range"), dict) else {}
            start_page = int(page_range.get("start_page", 0) or 0)
            end_page = int(page_range.get("end_page", 0) or 0)
            evidence_rows.append(
                f"[证据{idx}] doc_id={doc_id}, page={start_page}-{end_page}, chunk_id={chunk_id}"
            )
        if evidence_rows:
            note = doc.add_paragraph("证据：" + "；".join(evidence_rows))
            note.style = "ClauseText"


def _normalize_table_data(table_data: Any) -> tuple[list[str], list[dict[str, str]]]:
    if not isinstance(table_data, list) or not table_data:
        raise ValueError("table_data must be a non-empty list")
    if not all(isinstance(item, dict) for item in table_data):
        raise ValueError("table_data rows must be objects")

    headers: list[str] = []
    normalized_rows: list[dict[str, str]] = []
    for row in table_data:
        row_dict = dict(row)
        for key in row_dict.keys():
            key_text = _validate_text(str(key))
            if key_text not in headers:
                headers.append(key_text)
        normalized_rows.append({str(k): _validate_text(str(v)) for k, v in row_dict.items()})
    return headers, normalized_rows


def _append_table_block(doc: Document, block: dict[str, Any]) -> None:
    headers, rows = _normalize_table_data(block.get("table_data"))
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    if "Table Grid" in {style.name for style in doc.styles}:
        table.style = "Table Grid"

    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, header in enumerate(headers):
            table.cell(row_idx, col_idx).text = row.get(header, "")


def _append_image_meta_block(doc: Document, block: dict[str, Any]) -> None:
    _ensure_paragraph_style(doc, "ClauseText", _PARAGRAPH_STYLES)
    name = _validate_text(str(block.get("name") or ""))
    caption = _validate_text(str(block.get("caption") or ""))
    file_ref = _validate_text(str(block.get("file_ref") or ""))
    paragraph = doc.add_paragraph(f"图片名称：{name}；图片说明：{caption}；文件引用：{file_ref}")
    paragraph.style = "ClauseText"


def _append_attachment_meta_block(doc: Document, block: dict[str, Any]) -> None:
    _ensure_paragraph_style(doc, "ClauseText", _PARAGRAPH_STYLES)
    name = _validate_text(str(block.get("name") or ""))
    description = _validate_text(str(block.get("description") or ""))
    file_ref = _validate_text(str(block.get("file_ref") or ""))
    paragraph = doc.add_paragraph(f"附件名称：{name}；附件说明：{description}；文件引用：{file_ref}")
    paragraph.style = "ClauseText"


def _append_structured_blocks(doc: Document, blocks: list[dict[str, Any]]) -> None:
    for block in blocks:
        block_type = str(block.get("type") or "").strip()
        if block_type == "heading":
            _append_heading_block(doc, block)
        elif block_type == "paragraph":
            _append_paragraph_block(doc, block)
        elif block_type == "table":
            _append_table_block(doc, block)
        elif block_type == "image_meta":
            _append_image_meta_block(doc, block)
        elif block_type == "attachment_meta":
            _append_attachment_meta_block(doc, block)
        else:
            raise ValueError(f"unsupported content type: {block_type}")


def _export_pdf_via_soffice(docx_path: Path) -> str:
    cmd = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        str(docx_path),
        "--outdir",
        str(docx_path.parent),
    ]
    result = subprocess.run(cmd, check=False, capture_output=True, text=True)
    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        raise RuntimeError(f"LibreOffice export failed: {stderr or 'unknown error'}")
    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError("LibreOffice export failed: pdf output not found")
    return str(pdf_path)


def render_word_structured(
    output_path: str,
    content: dict[str, list[dict[str, Any]]],
    placeholders: dict[str, str] | None = None,
    template_path: str | None = None,
    style_config: dict | None = None,
    export_pdf: bool = False,
) -> tuple[str, str | None]:
    template_root = Path(settings.render_template_dir)
    export_root = Path(settings.render_output_dir)
    template_root.mkdir(parents=True, exist_ok=True)
    export_root.mkdir(parents=True, exist_ok=True)

    template_candidate = template_path or "default_tender_template.docx"
    template_file = _safe_path(template_candidate, template_root)
    ensure_default_template(template_file)

    out = _safe_path(output_path, export_root)
    if out.suffix.lower() != ".docx":
        raise ValueError("output_path must end with .docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    tpl = DocxTemplate(str(template_file))
    tpl.render(placeholders or {})
    tpl.save(str(out))

    body_blocks = content.get("body", [])
    appendix_blocks = content.get("appendix", [])
    if not isinstance(body_blocks, list) or not isinstance(appendix_blocks, list):
        raise ValueError("content.body and content.appendix must be arrays")

    doc = Document(str(out))
    _append_structured_blocks(doc, body_blocks)
    if appendix_blocks:
        _append_heading_block(doc, {"type": "heading", "style": "Title1", "text": "附件"})
        _append_structured_blocks(doc, appendix_blocks)

    _apply_page_layout(doc, style_config or {})
    doc.save(str(out))

    pdf_path = _export_pdf_via_soffice(out) if export_pdf else None
    return str(out), pdf_path


def render_word_sections(
    output_path: str,
    sections: list[dict[str, str]],
    template_path: str | None = None,
    title: str = "投标文件",
) -> str:
    template_root = Path(settings.render_template_dir)
    export_root = Path(settings.render_output_dir)
    template_root.mkdir(parents=True, exist_ok=True)
    export_root.mkdir(parents=True, exist_ok=True)

    template_candidate = template_path or "default_tender_template.docx"
    template_file = _safe_path(template_candidate, template_root)
    ensure_default_template(template_file)

    out = _safe_path(output_path, export_root)
    if out.suffix.lower() != ".docx":
        raise ValueError("output_path must end with .docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_file))
    doc.add_heading(title, level=1)
    for idx, section in enumerate(sections, start=1):
        heading = section.get("title") or f"章节 {idx}"
        content = section.get("content") or ""
        doc.add_heading(heading, level=2)
        doc.add_paragraph(content)
        # Auto-insert signature placeholders for sections that need signing
        if any(sig_key in heading for sig_key in ("投标函", "授权委托书", "承诺函")):
            for label in ("授权代表签字", "投标人盖章", "日期"):
                _add_signature_placeholder(doc, label)

    doc.save(str(out))
    return str(out)
