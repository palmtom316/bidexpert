from __future__ import annotations

import io
import hashlib
import json
import re
import uuid
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.exc import SQLAlchemyError

from app.core.config import settings
from app.db.session import session_scope
from app.models.tables import DocKind, Document, EvidenceChunk, ExpertDoc, Project, SensitivityLevel
from app.schemas.contracts import (
    DocBlockItem,
    ExpertLibraryChunkItem,
    ExpertLibraryDocItem,
    ExpertLibraryIngestResponse,
    ExpertLibraryStructuredIngestItem,
    ExpertLibraryStructuredIngestResponse,
    EvidenceUpsertItem,
)
from app.services.expert_enterprise_pipeline import (
    build_exceptions_queue,
    build_structure_v1_from_blocks,
    chunks_for_enterprise_rag,
    enrich_sections_v1,
    merge_structure_meta_risk,
    render_enterprise_markdown,
    risk_review_sections,
    serialize_chunks_jsonl,
    summarize_tables_in_structure,
)
from app.services.expert_workspace import (
    ExpertDocWorkspace,
    ensure_expert_library_layout,
    prepare_doc_workspace,
    sync_enterprise_config_assets,
)
from app.services.pdf_ingest import build_doc_blocks, extract_pages
from app.services.pricing_guard import detect_pricing_content
from app.services.qdrant_store import get_qdrant_store

_STRUCTURED_CATEGORY_MAP = {
    "STANDARD": ("规范", "STANDARD_SPEC", "STANDARD"),
    "COMPANY_PERFORMANCE": ("公司业绩", "COMPANY_PERFORMANCE", "PERFORMANCE"),
    "COMPANY_QUALIFICATION": ("公司资质", "COMPANY_QUALIFICATION", "QUALIFICATION"),
    "PM_QUALIFICATION_PERFORMANCE": ("项目管理人员资质及业绩", "PM_QUAL_PERFORMANCE", "PM_TEAM"),
}


@dataclass
class _ParsedProject:
    project_uuid: uuid.UUID | None
    project_raw: str | None


def _parse_project_id(project_id: str | None) -> _ParsedProject:
    if not project_id:
        return _ParsedProject(project_uuid=None, project_raw=None)
    try:
        return _ParsedProject(project_uuid=uuid.UUID(project_id), project_raw=project_id)
    except ValueError:
        return _ParsedProject(project_uuid=None, project_raw=project_id)


def _ensure_project_exists(project_uuid: uuid.UUID | None) -> None:
    if not project_uuid:
        return
    with session_scope() as db:
        exists = db.execute(select(Project.id).where(Project.id == project_uuid)).scalar_one_or_none()
        if exists:
            return
        db.add(
            Project(
                id=project_uuid,
                name=f"Imported-{str(project_uuid)[:8]}",
                owner_user_id="system",
                description="Auto-created for expert library import",
            )
        )
        db.commit()


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "_", name).strip("_")
    return cleaned or "upload_file"


def _save_uploaded_file(filename: str, content: bytes, raw_bid_dir: Path) -> Path:
    target = raw_bid_dir / _safe_filename(filename)
    target.write_bytes(content)
    return target


def _save_json(path: Path, payload: dict | list) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _save_jsonl(path: Path, rows: list[dict]) -> None:
    content = "\n".join(json.dumps(row, ensure_ascii=False) for row in rows)
    path.write_text(f"{content}\n" if content else "", encoding="utf-8")


def _append_jsonl(path: Path, rows: list[dict]) -> None:
    if not rows:
        return
    with path.open("a", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False))
            f.write("\n")


def _structure_doc_type(doc_type: str) -> str:
    normalized = (doc_type or "").strip().upper()
    if "SPEC" in normalized:
        return "spec"
    if "MANUAL" in normalized:
        return "manual"
    if normalized:
        return "bid"
    return "other"


def _write_extracted_blocks(structure: dict, blocks_dir: Path) -> None:
    for section in structure.get("sections", []):
        for block in section.get("blocks", []):
            block_id = str(block.get("block_id") or "block")
            block_type = str(block.get("type", "")).lower()
            if block_type == "table":
                payload = block.get("table") or {}
                (blocks_dir / f"{block_id}.json").write_text(
                    json.dumps(payload, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            elif block_type == "text":
                (blocks_dir / f"{block_id}.txt").write_text(str(block.get("text", "")), encoding="utf-8")


def _write_review_yaml(path: Path, doc_id: str, exceptions: list[dict]) -> None:
    lines = [
        f"doc_id: {doc_id}",
        f"issue_count: {len(exceptions)}",
        "issues:",
    ]
    for item in exceptions:
        lines.append(f"  - issue: {item.get('issue')}")
        lines.append(f"    section_id: {item.get('section_id')}")
        lines.append(f"    action: {item.get('action')}")
        lines.append(f"    detail: {item.get('detail')}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _write_pipeline_run_log(
    *,
    layout,
    doc_id: str,
    filename: str,
    chunk_count: int,
    exception_count: int,
    warnings: list[str],
    pricing_blocked: bool,
) -> None:
    run_dir = layout.stage_dirs["99_logs"] / "pipeline_runs"
    run_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    payload = {
        "run_id": f"{stamp}_run01",
        "doc_id": doc_id,
        "filename": filename,
        "status": "SUCCEEDED",
        "chunk_count": chunk_count,
        "exception_count": exception_count,
        "pricing_blocked": pricing_blocked,
        "warnings": warnings,
        "stages": [
            "classify_files",
            "extract_structure",
            "summarize_tables",
            "enrich_sections",
            "risk_review",
            "merge_and_validate",
            "render_markdown",
            "chunk_for_rag",
        ],
    }
    _save_json(run_dir / f"{stamp}_run01.json", payload)


def _load_thresholds(layout) -> dict[str, float]:
    defaults: dict[str, float] = {
        "low_confidence": 0.60,
        "strong_review_confidence": 0.75,
        "max_section_pages": 20.0,
        "max_chunk_tokens": float(settings.expert_chunk_max_tokens),
        "chunk_overlap_tokens": float(settings.expert_chunk_overlap_tokens),
    }
    config_path = layout.root / "00_config" / "pipeline" / "thresholds.v1.yaml"
    if not config_path.exists():
        return defaults

    for raw_line in config_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or ":" not in line:
            continue
        key, value = line.split(":", 1)
        key = key.strip()
        value = value.strip()
        if key not in defaults:
            continue
        try:
            defaults[key] = float(value)
        except ValueError:
            continue
    return defaults


def _fallback_chunks_from_blocks(
    *,
    blocks: list,
    industry_tag: str | None,
    doc_type: str,
    pricing_related: bool,
    doc_id: str,
) -> list[EvidenceUpsertItem]:
    chunks: list[EvidenceUpsertItem] = []
    for idx, block in enumerate(blocks, start=1):
        text = (block.content_text or "").strip()
        if len(text) < 24:
            continue
        page_no = int(getattr(block, "page_no", 1) or 1)
        section_id = f"fb-sec-{idx:04d}"
        section_type = "TABLE" if str(getattr(block, "block_type", "")).upper() == "TABLE" else "PARA"
        digest = hashlib.sha1(text.encode("utf-8")).hexdigest()[:12]
        chunks.append(
            EvidenceUpsertItem(
                chunk_id=f"fb-{idx}-{digest}",
                text=text,
                doc_type=doc_type,
                section_type=section_type,
                industry_tag=industry_tag,
                sensitivity_level="PUBLIC_OK",
                forbidden_tags=["PRICING_RELATED"] if pricing_related else [],
                source_locator={
                    "doc_id": doc_id,
                    "section_id": section_id,
                    "section_type": section_type,
                    "discipline": industry_tag or "GENERAL",
                    "source_page": page_no,
                    "section_anchor": block.section_anchor,
                    "block_type": "table" if section_type == "TABLE" else "text",
                },
            )
        )
    return chunks


def _decode_text_content(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _looks_like_markdown_table(lines: list[str]) -> bool:
    if len(lines) < 2:
        return False
    if "|" not in lines[0] or "|" not in lines[1]:
        return False
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", lines[1]))


def _build_text_blocks(*, text: str, default_anchor: str, markdown_mode: bool = False) -> list[DocBlockItem]:
    blocks: list[DocBlockItem] = []
    current_anchor = (default_anchor or "未命名章节").strip()[:48] or "未命名章节"
    cursor = 0

    for paragraph in [item.strip() for item in re.split(r"\n{2,}", text) if item.strip()]:
        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        if not lines:
            continue
        first_line = lines[0]
        heading_match = re.match(r"^\s{0,3}#{1,6}\s*(.+)$", first_line)
        if heading_match:
            current_anchor = heading_match.group(1).strip()[:48] or current_anchor
        elif re.match(r"^\s*(第[一二三四五六七八九十0-9]+[章节条款]|\d+(?:\.\d+)+)", first_line):
            current_anchor = first_line[:48]

        block_type = "TABLE" if markdown_mode and _looks_like_markdown_table(lines) else "PARA"
        start = cursor
        end = cursor + len(paragraph)
        cursor = end + 1
        blocks.append(
            DocBlockItem(
                page_no=1,
                block_type=block_type,
                section_anchor=current_anchor,
                content_text=paragraph,
                char_start=start,
                char_end=end,
            )
        )
    return blocks


def _extract_docx_blocks(filename: str, content: bytes) -> list[DocBlockItem]:
    from docx import Document as WordDocument

    try:
        doc = WordDocument(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise ValueError("invalid .docx content") from exc

    blocks: list[DocBlockItem] = []
    current_anchor = (Path(filename).stem or "未命名章节").strip()[:48] or "未命名章节"
    cursor = 0

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "").lower()
        if "heading" in style_name:
            current_anchor = text[:48] or current_anchor
            continue
        start = cursor
        end = cursor + len(text)
        cursor = end + 1
        blocks.append(
            DocBlockItem(
                page_no=1,
                block_type="PARA",
                section_anchor=current_anchor,
                content_text=text,
                char_start=start,
                char_end=end,
            )
        )

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        table_text = "\n".join(" | ".join(cell or "-" for cell in row) for row in rows)
        start = cursor
        end = cursor + len(table_text)
        cursor = end + 1
        blocks.append(
            DocBlockItem(
                page_no=1,
                block_type="TABLE",
                section_anchor=current_anchor,
                content_text=table_text,
                char_start=start,
                char_end=end,
            )
        )

    return blocks


def _extract_upload_blocks(
    filename: str,
    content: bytes,
) -> tuple[list[DocBlockItem], int, str, str, str]:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        pages = extract_pages(content, enable_ocr_fallback=settings.enable_ocr_fallback)
        return (
            build_doc_blocks(pages),
            len(pages),
            "pdf",
            "application/pdf",
            "pdf_ingest.v1",
        )
    if ext in {".md", ".markdown"}:
        text = _decode_text_content(content)
        blocks = _build_text_blocks(text=text, default_anchor=Path(filename).stem, markdown_mode=True)
        if not blocks:
            raise ValueError("empty markdown content")
        return (
            blocks,
            1,
            "markdown",
            "text/markdown",
            "markdown_ingest.v1",
        )
    if ext == ".docx":
        blocks = _extract_docx_blocks(filename, content)
        if not blocks:
            raise ValueError("empty word document content")
        return (
            blocks,
            1,
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx_ingest.v1",
        )
    if ext == ".doc":
        raise ValueError("暂不支持 .doc，请另存为 .docx 后上传")
    raise ValueError("unsupported file format, allowed: .pdf/.md/.markdown/.docx")


def ingest_historical_pdf(
    *,
    filename: str,
    content: bytes,
    project_id: str | None,
    industry_tag: str | None,
    title: str | None,
    created_by: str = "system",
    doc_type: str = "EXPERT_HISTORY",
    model_id: str | None = None,
) -> ExpertLibraryIngestResponse:
    parsed_project = _parse_project_id(project_id)
    _ensure_project_exists(parsed_project.project_uuid)

    layout = ensure_expert_library_layout()
    sync_enterprise_config_assets(layout)
    thresholds = _load_thresholds(layout)

    blocks, page_count, source_format, content_type, parser_version = _extract_upload_blocks(filename, content)
    full_text = "\n".join((block.content_text or "") for block in blocks if block.content_text)

    pricing_blocked, pricing_reasons = detect_pricing_content(full_text)
    warnings: list[str] = []
    if (model_id or "").strip():
        warnings.append("model_id_ignored_in_section_chunking_pipeline")
    if pricing_blocked:
        warnings.extend([f"pricing_detected:{reason}" for reason in pricing_reasons])
        warnings.append("section_enhancement_chunking_with_pricing_redaction_flags")

    doc_workspace: ExpertDocWorkspace | None = None
    structure: dict[str, Any] | None = None
    table_summaries: list[dict] = []
    section_meta: list[dict] = []
    risk_reviews: list[dict] = []
    merged: dict[str, Any] | None = None
    exception_queue: list[dict] = []
    markdown: str | None = None

    source_document_id: uuid.UUID | None = None
    expert_doc_id: uuid.UUID
    chunks: list[EvidenceUpsertItem] = []

    try:
        with session_scope() as db:
            doc_workspace = prepare_doc_workspace(layout, title or Path(filename).stem)
            saved_path = _save_uploaded_file(filename, content, doc_workspace.raw_bid_dir)

            source_doc = Document(
                project_id=parsed_project.project_uuid,
                kind=DocKind.EXPERT,
                filename=filename,
                content_type=content_type,
                object_uri=str(saved_path),
                sha256=hashlib.sha256(content).hexdigest(),
                page_count=page_count,
                language="zh-CN",
                sensitivity=SensitivityLevel.PUBLIC_OK,
                created_by=created_by,
            )
            db.add(source_doc)
            db.flush()
            source_document_id = source_doc.id
            doc_id = f"{doc_workspace.doc_key}-{str(source_document_id)[:8]}"

            structure = build_structure_v1_from_blocks(
                doc_id=doc_id,
                title=title,
                source_file=filename,
                source_format=source_format,
                blocks=blocks,
                parser_version=parser_version,
                doc_type=_structure_doc_type(doc_type),
            )
            table_summaries = summarize_tables_in_structure(structure)
            section_meta = enrich_sections_v1(structure, table_summaries)
            risk_reviews = risk_review_sections(
                structure,
                section_meta,
                strong_review_confidence=float(thresholds["strong_review_confidence"]),
            )
            merged = merge_structure_meta_risk(structure, section_meta, risk_reviews)
            exception_queue = build_exceptions_queue(
                doc_id=doc_id,
                merged=merged,
                low_confidence=float(thresholds["low_confidence"]),
                max_section_pages=int(thresholds["max_section_pages"]),
            )
            markdown = render_enterprise_markdown(merged)
            chunks = chunks_for_enterprise_rag(
                merged,
                industry_tag=industry_tag,
                doc_type=doc_type,
                min_tokens=settings.expert_chunk_min_tokens,
                max_tokens=int(thresholds["max_chunk_tokens"]),
                overlap_tokens=int(thresholds["chunk_overlap_tokens"]),
            )

            if not chunks:
                chunks = _fallback_chunks_from_blocks(
                    blocks=blocks,
                    industry_tag=industry_tag,
                    doc_type=doc_type,
                    pricing_related=pricing_blocked,
                    doc_id=doc_id,
                )
                warnings.append("no_section_chunks_fallback_to_block_chunks")

            if pricing_blocked:
                for chunk in chunks:
                    tags = list(chunk.forbidden_tags or [])
                    if "PRICING_RELATED" not in tags:
                        tags.append("PRICING_RELATED")
                    chunk.forbidden_tags = tags

            expert_doc = ExpertDoc(
                source_document_id=source_document_id,
                doc_type=doc_type,
                title=title or filename,
                industry_tag=industry_tag,
                section_type="BID_HISTORY",
                sensitivity=SensitivityLevel.PUBLIC_OK,
                valid_from=date.today(),
                created_by=created_by,
            )
            db.add(expert_doc)
            db.flush()
            expert_doc_id = expert_doc.id

            for idx, chunk in enumerate(chunks, start=1):
                source_locator = chunk.source_locator or {}
                source_page = source_locator.get("source_page")
                db.add(
                    EvidenceChunk(
                        expert_doc_id=expert_doc_id,
                        chunk_no=idx,
                        excerpt_text=chunk.text,
                        excerpt_hash=hashlib.sha1(chunk.text.encode("utf-8")).hexdigest(),
                        location={
                            "section_anchor": source_locator.get("section_title")
                            or source_locator.get("section_anchor"),
                            "page_no": source_page,
                        },
                        source_locator=source_locator,
                        valid_to=date.fromisoformat(chunk.valid_to) if chunk.valid_to else None,
                        sensitivity_level=SensitivityLevel.PUBLIC_OK,
                        quality_score=float(chunk.quality_score),
                        forbidden_tags=chunk.forbidden_tags or [],
                        qdrant_point_id=str(uuid.uuid5(uuid.NAMESPACE_URL, f"{expert_doc_id}:{chunk.chunk_id}")),
                        parent_chunk_id=chunk.parent_chunk_id or source_locator.get("parent_chunk_id"),
                        anchor_type=chunk.anchor_type or source_locator.get("anchor_type"),
                    )
                )

            db.commit()
    except SQLAlchemyError as exc:
        raise RuntimeError(f"failed to persist expert library records: {exc}") from exc

    store = get_qdrant_store()
    upserted = store.upsert_chunks(str(expert_doc_id), chunks, project_id=parsed_project.project_raw)

    if source_document_id and doc_workspace and structure and merged and markdown is not None:
        _save_json(doc_workspace.extracted_dir / "structure.v1.json", structure)
        _write_extracted_blocks(structure, doc_workspace.extracted_blocks_dir)
        _save_jsonl(doc_workspace.enriched_dir / "table_summary.v1.jsonl", table_summaries)
        _save_jsonl(doc_workspace.enriched_dir / "section_meta.v1.jsonl", section_meta)
        _save_jsonl(doc_workspace.enriched_dir / "risk_review.v1.jsonl", risk_reviews)
        _save_json(doc_workspace.enriched_dir / "merged.v1.json", merged)

        (layout.stage_dirs["04_md"] / f"{doc_workspace.doc_key}.enhanced.md").write_text(markdown, encoding="utf-8")
        _save_jsonl(doc_workspace.chunks_dir / "chunks.v1.jsonl", serialize_chunks_jsonl(chunks))

        _save_json(
            layout.stage_dirs["06_index"] / "qdrant" / f"{doc_workspace.doc_key}.json",
            {
                "expert_doc_id": str(expert_doc_id),
                "source_document_id": str(source_document_id),
                "qdrant_upserted": upserted,
                "chunk_count": len(chunks),
            },
        )
        _append_jsonl(layout.stage_dirs["07_review"] / "exceptions.queue.jsonl", exception_queue)
        _write_review_yaml(
            layout.stage_dirs["07_review"] / f"{doc_workspace.doc_key}.review.yaml",
            doc_id=str(merged.get("doc_id", "")),
            exceptions=exception_queue,
        )
        _write_pipeline_run_log(
            layout=layout,
            doc_id=str(merged.get("doc_id", "")),
            filename=filename,
            chunk_count=len(chunks),
            exception_count=len(exception_queue),
            warnings=warnings,
            pricing_blocked=pricing_blocked,
        )

    return ExpertLibraryIngestResponse(
        status="SUCCEEDED",
        expert_doc_id=str(expert_doc_id),
        source_document_id=str(source_document_id) if source_document_id else None,
        filename=filename,
        page_count=page_count,
        chunk_count=len(chunks),
        qdrant_upserted=upserted,
        warnings=warnings,
    )


def list_expert_docs(project_id: str | None, industry_tag: str | None, limit: int = 50) -> list[ExpertLibraryDocItem]:
    parsed_project = _parse_project_id(project_id)
    with session_scope() as db:
        stmt = (
            select(
                ExpertDoc.id,
                ExpertDoc.title,
                ExpertDoc.industry_tag,
                ExpertDoc.doc_type,
                ExpertDoc.created_at,
                func.count(EvidenceChunk.id),
            )
            .outerjoin(EvidenceChunk, EvidenceChunk.expert_doc_id == ExpertDoc.id)
            .group_by(ExpertDoc.id)
            .order_by(ExpertDoc.created_at.desc())
            .limit(max(1, min(limit, 200)))
        )
        if industry_tag:
            stmt = stmt.where(ExpertDoc.industry_tag == industry_tag)
        if parsed_project.project_uuid:
            stmt = stmt.join(Document, Document.id == ExpertDoc.source_document_id).where(
                Document.project_id == parsed_project.project_uuid
            )

        rows = db.execute(stmt).all()
        return [
            ExpertLibraryDocItem(
                expert_doc_id=str(row[0]),
                title=row[1],
                industry_tag=row[2],
                doc_type=row[3],
                created_at=row[4].isoformat() if isinstance(row[4], datetime) else str(row[4]),
                chunk_count=int(row[5] or 0),
            )
            for row in rows
        ]


def list_expert_chunks(expert_doc_id: str, limit: int = 200) -> list[ExpertLibraryChunkItem]:
    try:
        doc_uuid = uuid.UUID(expert_doc_id)
    except ValueError as exc:
        raise ValueError("invalid expert_doc_id") from exc

    with session_scope() as db:
        stmt = (
            select(EvidenceChunk)
            .where(EvidenceChunk.expert_doc_id == doc_uuid)
            .order_by(EvidenceChunk.chunk_no.asc())
            .limit(max(1, min(limit, 1000)))
        )
        rows = db.execute(stmt).scalars().all()
        items: list[ExpertLibraryChunkItem] = []
        for row in rows:
            source_locator = row.source_locator or {}
            items.append(
                ExpertLibraryChunkItem(
                    chunk_id=row.excerpt_hash or str(row.chunk_no),
                    excerpt_text=row.excerpt_text,
                    section_anchor=source_locator.get("section_anchor"),
                    quality_score=float(row.quality_score or 0),
                    valid_to=row.valid_to.isoformat() if row.valid_to else None,
                    created_at=row.created_at.isoformat() if row.created_at else "",
                )
            )
        return items


def _clean_lines(items: list[str]) -> list[str]:
    return [line.strip() for line in items if line and line.strip()]


def _build_structured_chunks(
    *,
    category_key: str,
    lines: list[str],
    industry_tag: str | None,
) -> list[EvidenceUpsertItem]:
    chunks: list[EvidenceUpsertItem] = []
    _, doc_type, section_type = _STRUCTURED_CATEGORY_MAP[category_key]
    for idx, text in enumerate(lines, start=1):
        digest = hashlib.sha1(f"{category_key}:{text}".encode("utf-8")).hexdigest()[:12]
        blocked, _ = detect_pricing_content(text)
        chunks.append(
            EvidenceUpsertItem(
                chunk_id=f"{category_key.lower()}-{idx}-{digest}",
                text=text,
                doc_type=doc_type,
                section_type=section_type,
                industry_tag=industry_tag,
                sensitivity_level="PUBLIC_OK",
                forbidden_tags=["PRICING_RELATED"] if blocked else [],
                quality_score=88.0,
                source_locator={"source": "structured_form", "category": category_key, "line": idx},
            )
        )
    return chunks


def _persist_structured_category(
    *,
    project_uuid: uuid.UUID | None,
    project_id: str | None,
    industry_tag: str | None,
    created_by: str,
    category_key: str,
    chunks: list[EvidenceUpsertItem],
) -> ExpertLibraryStructuredIngestItem:
    title, doc_type, section_type = _STRUCTURED_CATEGORY_MAP[category_key]
    text_blob = "\n".join(chunk.text for chunk in chunks)
    warnings: list[str] = []
    pricing_blocked, reasons = detect_pricing_content(text_blob)
    if pricing_blocked:
        warnings.extend([f"pricing_detected:{reason}" for reason in reasons])

    try:
        with session_scope() as db:
            source_doc = Document(
                project_id=project_uuid,
                kind=DocKind.EXPERT,
                filename=f"{title}.txt",
                content_type="text/plain",
                object_uri=f"local://expert_library/structured/{uuid.uuid4()}",
                sha256=hashlib.sha256(text_blob.encode("utf-8")).hexdigest(),
                page_count=1,
                language="zh-CN",
                sensitivity=SensitivityLevel.PUBLIC_OK,
                created_by=created_by,
            )
            db.add(source_doc)
            db.flush()

            expert_doc = ExpertDoc(
                source_document_id=source_doc.id,
                doc_type=doc_type,
                title=title,
                industry_tag=industry_tag,
                section_type=section_type,
                sensitivity=SensitivityLevel.PUBLIC_OK,
                valid_from=date.today(),
                created_by=created_by,
            )
            db.add(expert_doc)
            db.flush()

            for idx, chunk in enumerate(chunks, start=1):
                source_locator = dict(chunk.source_locator or {})
                source_locator.setdefault("doc_id", str(source_doc.id))
                source_locator.setdefault("section_id", f"{category_key.lower()}-{idx:03d}")
                source_locator.setdefault("section_type", section_type)
                source_locator.setdefault("discipline", industry_tag or "GENERAL")
                source_locator.setdefault("source_page", 1)
                source_locator.setdefault("block_type", "text")
                source_locator.setdefault("section_title", title)
                chunk.source_locator = source_locator
                db.add(
                    EvidenceChunk(
                        expert_doc_id=expert_doc.id,
                        chunk_no=idx,
                        excerpt_text=chunk.text,
                        excerpt_hash=hashlib.sha1(chunk.text.encode("utf-8")).hexdigest(),
                        location={"section_anchor": section_type, "page_no": 1},
                        source_locator=source_locator,
                        valid_to=date.fromisoformat(chunk.valid_to) if chunk.valid_to else None,
                        sensitivity_level=SensitivityLevel.PUBLIC_OK,
                        quality_score=float(chunk.quality_score),
                        forbidden_tags=chunk.forbidden_tags or [],
                        qdrant_point_id=str(uuid.uuid5(uuid.NAMESPACE_URL, f"{expert_doc.id}:{chunk.chunk_id}")),
                        parent_chunk_id=chunk.parent_chunk_id or source_locator.get("parent_chunk_id"),
                        anchor_type=chunk.anchor_type or source_locator.get("anchor_type"),
                    )
                )

            db.commit()
            expert_doc_id = str(expert_doc.id)
    except SQLAlchemyError as exc:
        raise RuntimeError(f"failed to persist structured expert library records: {exc}") from exc

    store = get_qdrant_store()
    upserted = store.upsert_chunks(expert_doc_id, chunks, project_id=project_id)
    return ExpertLibraryStructuredIngestItem(
        category=category_key,
        expert_doc_id=expert_doc_id,
        title=title,
        chunk_count=len(chunks),
        qdrant_upserted=upserted,
        warnings=warnings,
    )


def ingest_structured_expert_knowledge(
    *,
    project_id: str | None,
    industry_tag: str | None,
    created_by: str,
    standard_items: list[str],
    company_performance_items: list[str],
    company_qualification_items: list[str],
    pm_qualification_performance_items: list[str],
) -> ExpertLibraryStructuredIngestResponse:
    parsed_project = _parse_project_id(project_id)
    _ensure_project_exists(parsed_project.project_uuid)

    grouped = {
        "STANDARD": _clean_lines(standard_items),
        "COMPANY_PERFORMANCE": _clean_lines(company_performance_items),
        "COMPANY_QUALIFICATION": _clean_lines(company_qualification_items),
        "PM_QUALIFICATION_PERFORMANCE": _clean_lines(pm_qualification_performance_items),
    }
    if not any(grouped.values()):
        raise ValueError("at least one structured item is required")

    items: list[ExpertLibraryStructuredIngestItem] = []
    total_chunks = 0
    for category_key, lines in grouped.items():
        if not lines:
            continue
        chunks = _build_structured_chunks(
            category_key=category_key,
            lines=lines,
            industry_tag=industry_tag,
        )
        item = _persist_structured_category(
            project_uuid=parsed_project.project_uuid,
            project_id=parsed_project.project_raw,
            industry_tag=industry_tag,
            created_by=created_by,
            category_key=category_key,
            chunks=chunks,
        )
        items.append(item)
        total_chunks += item.chunk_count

    return ExpertLibraryStructuredIngestResponse(
        status="SUCCEEDED",
        total_docs=len(items),
        total_chunks=total_chunks,
        items=items,
    )
