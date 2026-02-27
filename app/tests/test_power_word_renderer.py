"""Power engineering word renderer template tests."""
from __future__ import annotations

import tempfile
from pathlib import Path

from app.services.word_renderer import (
    ensure_default_template,
    _SIGNATURE_SECTIONS,
    _add_signature_placeholder,
)


def test_signature_sections_set_exists():
    assert len(_SIGNATURE_SECTIONS) >= 5
    assert "投标人盖章" in _SIGNATURE_SECTIONS
    assert "项目经理签章" in _SIGNATURE_SECTIONS


def test_ensure_default_template_creates_comprehensive_doc():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "test_template.docx"
        ensure_default_template(path)
        assert path.exists()
        # File should be substantially larger than a minimal template
        assert path.stat().st_size > 5000


def test_ensure_default_template_idempotent():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "test_template.docx"
        ensure_default_template(path)
        size1 = path.stat().st_size
        ensure_default_template(path)
        size2 = path.stat().st_size
        assert size1 == size2


def test_ensure_default_template_has_standard_sections():
    with tempfile.TemporaryDirectory() as tmpdir:
        from docx import Document

        path = Path(tmpdir) / "test_template.docx"
        ensure_default_template(path)
        doc = Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "投标文件" in full_text
        assert "技术方案" in full_text
        assert "施工组织设计" in full_text
        assert "质量保证方案" in full_text
        assert "安全生产方案" in full_text


def test_add_signature_placeholder():
    from docx import Document

    doc = Document()
    _add_signature_placeholder(doc, "投标人盖章")
    assert "投标人盖章" in doc.paragraphs[-1].text
