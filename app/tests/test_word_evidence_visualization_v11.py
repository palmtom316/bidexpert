from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX

from app.core.config import settings
from app.services.word_renderer import render_word_structured


def _prepare_template(tmp_path: Path, monkeypatch) -> str:
    template_dir = tmp_path / "templates"
    output_dir = tmp_path / "exports"
    template_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(settings, "render_template_dir", str(template_dir))
    monkeypatch.setattr(settings, "render_output_dir", str(output_dir))

    doc = Document()
    for style_name in ["Title1", "Title2", "Title3", "Title4", "BodyText", "BodyText_Indent", "ClauseText"]:
        if style_name not in [s.name for s in doc.styles]:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("项目名称：{{project_name}}")
    name = "template-v11.docx"
    doc.save(str(template_dir / name))
    return name


def test_structured_render_adds_evidence_note_and_highlights_high_risk(monkeypatch, tmp_path: Path) -> None:
    template_name = _prepare_template(tmp_path, monkeypatch)

    output_path, _ = render_word_structured(
        output_path="v11-evidence.docx",
        template_path=template_name,
        placeholders={"project_name": "示例项目"},
        content={
            "body": [
                {"type": "heading", "style": "Title1", "text": "章节A"},
                {
                    "type": "paragraph",
                    "style": "BodyText",
                    "text": "这是高风险段落。",
                    "risk_level": "high",
                    "evidence": [
                        {
                            "doc_id": "doc-1",
                            "page_range": {"start_page": 2, "end_page": 3},
                            "chunk_id": "chunk-1",
                        }
                    ],
                },
            ],
            "appendix": [],
        },
        style_config={},
        export_pdf=False,
    )

    doc = Document(output_path)
    paragraph = next(p for p in doc.paragraphs if p.text.strip() == "这是高风险段落。")
    assert paragraph.runs
    assert paragraph.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW

    evidence_note = [p.text for p in doc.paragraphs if "证据" in p.text and "chunk-1" in p.text]
    assert evidence_note
