from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app
from app.services.word_renderer import render_word_structured


def _prepare_template(monkeypatch, tmp_path: Path, template_name: str = "structured-template.docx") -> str:
    template_dir = tmp_path / "templates"
    output_dir = tmp_path / "exports"
    template_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(settings, "render_template_dir", str(template_dir))
    monkeypatch.setattr(settings, "render_output_dir", str(output_dir))

    doc = Document()
    for style_name in ["Title1", "Title2", "Title3", "Title4", "BodyText", "BodyText_Indent", "ClauseText"]:
        if style_name not in [s.name for s in doc.styles]:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("项目名称：{{project_name}}")
    doc.save(str(template_dir / template_name))
    return template_name


def test_structured_render_rejects_manual_numbered_title(monkeypatch, tmp_path: Path) -> None:
    template_name = _prepare_template(monkeypatch, tmp_path)

    try:
        render_word_structured(
            output_path="invalid.docx",
            template_path=template_name,
            placeholders={"project_name": "示例项目"},
            content={
                "body": [
                    {"type": "heading", "style": "Title1", "text": "1. 项目总体说明"},
                ],
                "appendix": [],
            },
            style_config={},
            export_pdf=False,
        )
        assert False, "expected ValueError for manual numbering in title"
    except ValueError as exc:
        assert "编号" in str(exc) or "number" in str(exc).lower()


def test_structured_render_applies_styles_and_table_data(monkeypatch, tmp_path: Path) -> None:
    template_name = _prepare_template(monkeypatch, tmp_path)

    output_path, pdf_path = render_word_structured(
        output_path="structured.docx",
        template_path=template_name,
        placeholders={"project_name": "光伏项目"},
        content={
            "body": [
                {"type": "heading", "style": "Title1", "text": "项目总体说明"},
                {"type": "paragraph", "style": "BodyText", "text": "这是正文段落。"},
                {
                    "type": "table",
                    "table_data": [
                        {"no": "1", "name": "技术负责人", "value": "张三"},
                        {"no": "2", "name": "项目经理", "value": "李四"},
                    ],
                },
            ],
            "appendix": [
                {"type": "paragraph", "style": "ClauseText", "text": "附件说明"},
                {"type": "image_meta", "name": "施工总平面图", "caption": "现场总图", "file_ref": "img-001"},
                {"type": "attachment_meta", "name": "营业执照", "description": "加盖公章扫描件", "file_ref": "att-001"},
            ],
        },
        style_config={},
        export_pdf=False,
    )

    assert output_path.endswith(".docx")
    assert pdf_path is None

    doc = Document(output_path)
    found = {p.text: p.style.name for p in doc.paragraphs if p.text.strip()}
    assert found["项目总体说明"] == "Title1"
    assert found["这是正文段落。"] in {"BodyText", "Body Text"}
    assert found["附件说明"] == "ClauseText"

    tables = doc.tables
    assert len(tables) >= 1
    assert tables[0].cell(0, 0).text == "no"
    assert tables[0].cell(1, 1).text == "技术负责人"
    assert tables[0].cell(2, 2).text == "李四"


def test_structured_render_api_rejects_forbidden_content(monkeypatch, tmp_path: Path) -> None:
    template_name = _prepare_template(monkeypatch, tmp_path)
    client = TestClient(app)

    response = client.post(
        "/v1/render/word/structured",
        json={
            "output_path": "api-invalid.docx",
            "template_path": template_name,
            "placeholders": {"project_name": "示例项目"},
            "content": {
                "body": [{"type": "heading", "style": "Title1", "text": "1.1 施工组织设计原则"}],
                "appendix": [],
            },
            "style_config": {},
            "export_pdf": False,
        },
    )
    assert response.status_code == 400


def test_structured_render_runs_soffice_for_pdf_export(monkeypatch, tmp_path: Path) -> None:
    template_name = _prepare_template(monkeypatch, tmp_path)
    called: dict[str, list[str]] = {}

    def fake_run(cmd, check, capture_output, text):  # noqa: ANN001
        called["cmd"] = cmd
        assert check is False
        assert capture_output is True
        assert text is True

        export_dir = Path(settings.render_output_dir)
        (export_dir / "deliverable.pdf").write_bytes(b"%PDF-1.4")

        class Result:
            returncode = 0
            stdout = "ok"
            stderr = ""

        return Result()

    monkeypatch.setattr("app.services.word_renderer.subprocess.run", fake_run)

    output_path, pdf_path = render_word_structured(
        output_path="deliverable.docx",
        template_path=template_name,
        placeholders={"project_name": "示例项目"},
        content={
            "body": [{"type": "heading", "style": "Title1", "text": "项目总体说明"}],
            "appendix": [],
        },
        style_config={},
        export_pdf=True,
    )

    assert output_path.endswith(".docx")
    assert pdf_path and pdf_path.endswith(".pdf")
    assert called["cmd"][0] == "soffice"
    assert "--headless" in called["cmd"]
    assert "--convert-to" in called["cmd"]
